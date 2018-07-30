import docx


def docx_to_text(file_path):
    print(file_path)
    doc = docx.Document(file_path)
    result = []
    tables = doc.tables
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    result.append(paragraph.text)

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt != '':
            txt = preprocess_text(txt)
            result.append(txt)
    return result


def preprocess_text(text):
    text = ' '.join(text.split())
    text = join_name_tag(text)
    return text


def join_name_tag(text):
    text = text.replace('\u2003', '')
    return text
